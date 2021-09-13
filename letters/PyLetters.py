from datetime import date
import datetime
import os.path
import logging
import pytz
from docx.shared import Pt
from docx import Document
import pyad.pyadutils
import pyad.adquery
import pyad.aduser
from letters import utils

# Todo Add logic to the password instructions function to account for the user being in Denver.

# Time/Date variables to get the current date for the logfile name.
now = datetime.datetime.now()
current_time = now.strftime("%H:%M:%S")
logging.basicConfig(filename=f'log_dir', level=logging.DEBUG)


def create_letters(q):
    """Class to gather required info and create the new hire letters"""
    # Variables used to help calculate and target users who have been created in the last 30 days
    today = datetime.datetime.utcnow().replace(tzinfo=pytz.UTC)
    days = datetime.timedelta(days=30)
    alpha = today - days

    # State which part of the program is running
    print("\n* * * * * * * * * * * * * * * * * * * * * * * *"
          "\n* * * * * * * * *  PyLetters  * * * * * * * * *"
          "\n* * * * * * * * * * * * * * * * * * * * * * * *\n")

    # Variable to store which users we created letters for - Used later in the email.
    users = []

    # Launching the AD query service, and then trying to use the service
    try:
        # q = pyad.adquery.ADQuery()
        logging.info(f"{current_time} - Gathering Active Directory Information...")
        # Giving the query service our search and return criteria
        q.execute_query(
            attributes=["whenCreated", "cn", "employeeID", "pwdLastSet", "description",
                        "displayName", "SamAccountName",
                        "mail",
                        "telephoneNumber",
                        "facsimileTelephoneNumber",
                        "Useraccountcontrol"],
            where_clause="objectClass = 'user'",
            base_dn="OU=Offices, DC=domain, DC=local"
        )
    except Exception as error:
        logging.error(f"{current_time} - {error}!")
        logging.error(f"{current_time} - Failed to initialize the Active Directory Query Service! "
                      f"Quitting!")
        return

    logging.info(f"{current_time} - Locating users created within the last two weeks...")
    # Checks to see if any users were pulled from the search, then filters out devices,
    # terminations, service accounts.
    # Also checks to see if the user pulled has been created in the last 30 days
    for row in q.get_results():
        # Makes sure we are not targeting devices (they end with $)
        if "$" not in str(row['displayName']):
            # This is checking to make sure the New Hire account is at maximum 30 days of age
            if row['whenCreated'] >= alpha:
                # This is checking to see if the user has a display name and if their telephone
                # is filled out
                if row['displayName'] is not None and row['telephoneNumber'] is not None:
                    # This determines whether the user's account is disabled or terminated.
                    if (not row["Useraccountcontrol"] / 2 % 2 != 0) or \
                            ("zz" not in str(row['displayName'])):
                        # ID 9999 is our ID for service accounts, we are filtering those out.
                        if row['employeeID'] != 9999:
                            file_path = "undelivered_dir"
                            delivered_path = "delivered_dir"

                            # Checks to see if the user has a fax number,
                            # chooses the template accordingly
                            if row['facsimileTelephoneNumber'] is not None:
                                doc = Document("Template_with_fax.docx")
                                doc2 = Document("Template_with_fax_manager.docx")
                                logging.info(
                                    f"{current_time} - Preparing {row['displayName']}'s letter with"
                                    f" Fax Fields...")
                            else:
                                doc = Document("Template.docx")
                                doc2 = Document("Template_manager.docx")
                                logging.info(
                                    f"{current_time} - Preparing {row['displayName']}'s letter "
                                    f"without Fax Fields...")

                            # Checks to see if the user already has a letter that has been delivered
                            if os.path.isdir(f"{delivered_path}\\{row['displayName']}"):
                                logging.warning(
                                    f"{current_time} - {row['displayName']} already has a folder "
                                    f"in the Delivered "
                                    f"Letters folder... Skipping user!")
                                print(f"{row['displayName']}'s folder exists in delivered "
                                      f"letters... Skipping...")
                                continue

                            # Checks to see if the letters exist,
                            # and if not creates each one even if one is missing
                            if os.path.isdir(f"{file_path}\\{row['displayName']}") and \
                                    os.path.isfile(f"{file_path}\\{row['displayName']}\\" +
                                                   f"New Hire Letter {row['displayName']}.docx") \
                                    and \
                                    os.path.isfile(f"{file_path}\\{row['displayName']}\\" +
                                                   f"New Hire Letter {row['displayName']} "
                                                   f"Manager.docx"):
                                logging.warning(
                                    f"{current_time} - {row['displayName']} already has a folder, "
                                    f"and 2 letters... "
                                    f"Skipping user!")
                                print(f"New Hire Letter {row['displayName']}.docx exists... "
                                      f"Skipping...")
                                continue
                            elif os.path.isdir(f"{file_path}\\{row['displayName']}") is False:
                                logging.info(f"{current_time} - "
                                             f"Creating {row['displayName']} folder...")
                                os.mkdir(f"{file_path}\\{row['displayName']}")
                            else:
                                logging.warning(
                                    f"{current_time} - {row['displayName']} already has a "
                                    f"directory and 1 letter, attempted to generate other letter. "
                                    f"Please review this new hire letter!")

                            # Takes the pwdLastSet and whenCreated attribute and converts them
                            # to a date that is
                            # readable for the program to do its calculations
                            last_set = pyad.pyadutils.convert_bigint(row['pwdLastSet'])
                            created = row['whenCreated'].strftime('%Y-%m-%d %H:%M:%S')
                            created = datetime.datetime.strptime(created, "%Y-%m-%d %H:%M:%S")
                            timestamp = last_set
                            last_set = datetime.datetime(1601, 1, 1) + \
                                       datetime.timedelta(seconds=timestamp / 10000000)
                            last_set = last_set.strftime('%Y-%m-%d %H:%M:%S')
                            last_set = datetime.datetime.strptime(last_set, "%Y-%m-%d %H:%M:%S")
                            elapsed_time = last_set - created

                            # This is checking to see if the user's Creation Time and Password
                            # Last Set are within 5
                            # seconds of each other and changes the password if they are.
                            if elapsed_time < datetime.timedelta(seconds=5):
                                password = utils.reset_user_password(row['cn'])
                                logging.info(f"{current_time} - {row['displayName']}'s password "
                                             f"has been changed.")
                            else:
                                print(f"{row['displayName']} does not need a password change\n")
                                logging.warning(
                                    f"{current_time} - {row['displayName']}'s password has NOT "
                                    f"been changed.")
                                password = "WARNING - COULD NOT SET PASSWORD"

                            # Initializing the styles we need for the docx
                            style = doc.styles['Normal']
                            font = style.font
                            font.name = 'Calibri'
                            font.size = Pt(11)

                            style2 = doc2.styles['Normal']
                            font2 = style2.font
                            font2.name = 'Calibri'
                            font2.size = Pt(11)

                            # From here to line 195 is assigning the values we need to certain
                            # fields so when its
                            # saved they will be in their proper locations.
                            description = str(row["description"])
                            for sym in (("'", ""), (",", ""), ("(", ""), (")", "")):
                                description = description.replace(*sym)
                            device_type = utils.find_job_title(description.replace(",", ""))

                            user_table = doc.tables[0]
                            user_table.cell(0, 1).text = row['SamAccountName']
                            user_table.cell(1, 1).text = password

                            man_table = doc2.tables[0]
                            man_table.cell(0, 1).text = row['SamAccountName']
                            man_table.cell(1, 1).text = password

                            user_table = doc.tables[1]
                            user_table.cell(0, 0).text = device_type

                            man_table = doc2.tables[1]
                            man_table.cell(0, 0).text = device_type

                            user_table = doc.tables[2]
                            user_table.cell(0, 1).text = row['mail']

                            man_table = doc2.tables[2]
                            man_table.cell(0, 1).text = row['mail']

                            user_table = doc.tables[3]
                            user_table.cell(0, 1).text = row['telephoneNumber']
                            if row['facsimileTelephoneNumber'] is not None:
                                user_table.cell(2, 1).text = row['facsimileTelephoneNumber']
                            else:
                                pass

                            man_table = doc2.tables[3]
                            man_table.cell(0, 1).text = row['telephoneNumber']
                            if row['facsimileTelephoneNumber'] is not None:
                                man_table.cell(2, 1).text = row['facsimileTelephoneNumber']
                            else:
                                pass

                            user_table = doc.tables[4]
                            user_table.cell(0, 1).text = row['mail']

                            man_table = doc2.tables[4]
                            man_table.cell(0, 1).text = row['mail']

                            user_table = doc.tables[5]
                            user_table.cell(1, 1).text = row['SamAccountName']

                            user_table = doc.tables[6]
                            user_table.cell(1, 1).text = row['SamAccountName']

                            user_table = doc.tables[7]
                            user_table.cell(0, 1).text = row['SamAccountName']

                            # Attempting to save the two documents
                            # based off of the values we assigned above.
                            try:
                                logging.info(f"{current_time} - Beginning User and Manager letter "
                                             f"generation...")
                                doc.save(
                                    f"{file_path}\\{row['displayName']}\\" +
                                    f"New Hire Letter {row['displayName']}.docx")
                                doc2.save(
                                    f"{file_path}\\{row['displayName']}\\" +
                                    f"New Hire Letter {row['displayName']} Manager.docx")
                                logging.info(
                                    f"{current_time} - Successfully created User and Manager "
                                    f"letter for "
                                    f"{row['displayName']}.")
                                users.append(row['displayName'])
                            except Exception as error:
                                logging.error(f"{current_time} - {error}!")
                                logging.error(f"{current_time} - Failed to create Letters for "
                                              f"{row['displayName']}!")
                        else:
                            logging.warning(f"Skipping {row['displayName']}, "
                                            f"detected service account.")
                            print(f"Skipping {row['displayName']}, detected service account.")
                            continue
                    else:
                        logging.warning(f"{current_time} - Skipping {row['displayName']}, "
                                        f"they are disabled.")
                        print(f"Skipping {row['displayName']}, they are disabled.\n")
                        continue
                else:
                    logging.info(f"Skipping {row['displayName']}, "
                                 f"their AD Profile is not yet filled out..")
                    print(f"{row['displayName']}'s AD Profile is not yet filled out..")
                    continue
            else:
                continue
        else:
            continue

    # Checks to see if any users had a letter generated and sends out the Log and results email.
    logging.info(f"{current_time} - End of current run.\n")
    logging.shutdown()
    if users:
        users.sort()
        utils.send_log(users)
    else:
        pass

    # Utility to clean up our logging folder.
    utils.log_cleanup()
